# -*- coding: utf-8 -*-
"""生成《新产品测试上位机软件操作使用规范》Word 文档。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(r"D:\C\new_production_h_develop\docs\新产品测试上位机软件操作使用规范（初稿）.docx")


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, *, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, first_line=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading_cn(doc, text, level=1):
    # 用段落模拟标题，避免中文字体在 Heading 样式里异常
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12 if level == 1 else 8)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, name="黑体", size=sizes.get(level, 12), bold=True)
    # 大纲级别，便于 Word 生成目录
    pPr = p._p.get_or_add_pPr()
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        from lxml import etree

        outline = etree.SubElement(pPr, qn("w:outlineLvl"))
    outline.set(qn("w:val"), str(level - 1))
    return p


def set_cell_text(cell, text, *, bold=False, size=10.5, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, center=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            set_cell_text(table.rows[r_i + 1].cells[c_i], str(val), size=10.5)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_steps(doc, steps):
    for i, s in enumerate(steps, 1):
        add_para(doc, f"{i}. {s}", first_line=False, space_after=4)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ---- 封面 ----
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "受控等级：内部", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "文件编号：待定", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "新产品测试上位机", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "软件操作使用规范", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
    add_para(doc, "版本号 / 修改号：A/0", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "2026 年 7 月 27 日", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "适用程序：new_production.exe", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "文档性质：初稿（UI 操作规范）", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ---- 前言 ----
    add_heading_cn(doc, "前言", 1)
    add_para(
        doc,
        "本规范规定新产品测试上位机（以下简称「上位机」）的软件操作使用方法，重点说明工站切换、"
        "自由工站日常测试、测试流程编排以及工站内部测试步骤（功能块）的配置方法。",
        first_line=True,
    )
    add_para(
        doc,
        "本规范由产测软件维护人员编制，供产线操作员、工艺工程师、设备维护人员使用。"
        "界面以实机菜单与对话框文字为准；若与本文不一致，以实机为准。",
        first_line=True,
    )
    add_para(doc, "本规范为首次发布。", first_line=True)

    # ---- 1 范围 ----
    add_heading_cn(doc, "1 范围", 1)
    add_para(doc, "本规范适用于 Windows 平台 new_production.exe 在下列场景中的操作：", first_line=True)
    add_para(doc, "a）切换上位机工站类型（SYSTEM/station）并重启进入对应工站窗口；")
    add_para(doc, "b）在自由工站中连接外设、扫码开测、查看结果与日志；")
    add_para(doc, "c）在「测试流程编排」中管理工站 Profile、编排主流程与失败执行区；")
    add_para(doc, "d）在「测试项配置」对话框中配置步骤的通信类型、指令、参数、Gate、Hook 等。")
    add_para(
        doc,
        "本规范不包含各产品固件协议细节、MES 后台配置、云端平台运维，以及专项工站内部状态机原理说明。",
        first_line=True,
    )

    # ---- 2 术语 ----
    add_heading_cn(doc, "2 术语和定义", 1)
    add_table(
        doc,
        ["术语", "定义"],
        [
            ["上位机", "本程序 new_production.exe，用于产线测试与设备控制。"],
            ["工站类型", "上位机当前打开的工站窗口类型，由「工站选择」写入 SYSTEM/station，切换后需重启。"],
            ["流程工站 / Profile", "自由工站内「测试流程编排」所选的产品工站配置（显示名如「Wellness Warm组装吸力测试工站」）。"],
            ["自由工站", "按可配置测试流程逐步执行的工站窗口，标题一般为「自由工站」。"],
            ["一拖多", "同一工站窗口内按「行×列」显示多个工位面板。"],
            ["功能块 / 步骤", "流程中的一个测试项，对应 test_case 下的步骤 ini。"],
            ["Gate", "根据设备回传结果判定本步通过或失败。"],
            ["Hook", "不走普通测试指令、走预置复杂流程的特殊步骤。"],
            ["VISA", "程控仪器资源访问接口（如 GPIB、TCPIP 程控电源）。"],
            ["sharedPair", "步骤参数：启用多工位共享同一台外设（电源、温度仪等）。"],
        ],
        col_widths=[4, 12],
    )

    # ---- 3 要求 ----
    add_heading_cn(doc, "3 要求", 1)
    add_heading_cn(doc, "3.1 人员要求", 2)
    add_table(
        doc,
        ["角色", "要求"],
        [
            ["操作员", "能连接串口、扫码开测、查看通过/失败与日志；一般不进入「功能设置」。"],
            ["工艺工程师", "具备管理员或工艺工程师账号；能切换工站、编排流程、配置步骤参数。"],
            ["维护人员", "能排查串口占用、VISA 连接、流程失败项与上位机日志。"],
        ],
        col_widths=[3.5, 12.5],
    )
    add_heading_cn(doc, "3.2 设备与环境要求", 2)
    add_para(doc, "a）PC 已安装本上位机及必要依赖（如使用程控电源时需 NI-VISA）；")
    add_para(doc, "b）按工序准备 Dongle、治具串口、产品串口、万用表、程控电源、温度仪等；")
    add_para(doc, "c）串口/GPIB 未被其它软件占用（使用 GPIB 前建议关闭 NI Communicator）；")
    add_para(doc, "d）需要 MES 时网络可达；扫码枪切换为英文输入。")

    # ---- 4 界面概述 ----
    add_heading_cn(doc, "4 软件界面概述", 1)
    add_heading_cn(doc, "4.1 启动与登录", 2)
    add_steps(
        doc,
        [
            "运行 new_production.exe。",
            "按提示完成云平台/账号登录（若启用）。",
            "状态栏可查看「云平台：已登录」或「未登录」。",
            "程序按当前工站类型打开对应窗口（如「自由工站」）。",
        ],
    )
    add_heading_cn(doc, "4.2 菜单栏", 2)
    add_table(
        doc,
        ["菜单项", "说明"],
        [
            ["功能设置", "打开「上位机设置」。仅管理员或工艺工程师可见。"],
            ["连接治具串口", "自由工站：打开治具串口连接窗口。"],
            ["帮助 → 切换账号 / 上传日志 / 检查更新...", "账号与维护相关功能。"],
        ],
        col_widths=[5, 11],
    )
    add_heading_cn(doc, "4.3 上位机设置主要页签", 2)
    add_table(
        doc,
        ["页签", "主要内容"],
        [
            ["关键设置", "工站选择、MES、一拖多行列、产品差异化开关等。"],
            ["测试流程编排", "流程工站 Profile、主流程、失败执行区、上传/下载用例。"],
            ["测试详细配置", "电流阈值、按键、射频 PER、三元组、PLC 等专项参数。"],
        ],
        col_widths=[4, 12],
    )
    add_heading_cn(doc, "4.4 自由工站主界面（摘要）", 2)
    add_para(
        doc,
        "自由工站首 Tab 含扫码框、串口区、停止测试、结果表；另有日志、图形展示、蓝牙绑定等 Tab。"
        "开测方式为扫描 SN 或 MAC 后回车，无单独「开始」大按钮。日常测试操作见第 7 章。",
        first_line=True,
    )

    # ---- 5 工站切换（重点） ----
    add_heading_cn(doc, "5 工站切换操作说明", 1)
    add_para(
        doc,
        "工站切换分为两类，请勿混淆：① 上位机工站类型切换（更换程序窗口，如从吸力工站切到自由工站）；"
        "② 自由工站内流程工站（Profile）切换（同一自由工站窗口下更换产品测试流程）。",
        first_line=True,
    )

    add_heading_cn(doc, "5.1 上位机工站类型切换（更换窗口）", 2)
    add_heading_cn(doc, "5.1.1 操作步骤", 3)
    add_steps(
        doc,
        [
            "使用具备权限的账号登录上位机。",
            "在工站窗口菜单栏点击「功能设置」，打开「上位机设置」。",
            "进入「关键设置」页签，找到分组「1、工站选择」。",
            "点选目标工站单选项（见 5.1.2 一览表）。",
            "软件将保存 SYSTEM/station 配置，并提示需要重启；确认后程序自动关闭并重新启动。",
            "重启后核对窗口标题是否为目标工站（例如「自由工站」「吸力测试」等）。",
        ],
    )
    add_para(
        doc,
        "注意：切换工站类型会重启程序，请先结束正在进行的测试，避免中途丢失结果。",
        first_line=True,
    )

    add_heading_cn(doc, "5.1.2 工站选择一览", 3)
    add_table(
        doc,
        ["界面选项", "说明（用途）"],
        [
            ["调试上位机", "产品测试工具/调试主窗口。"],
            ["电机校准上位机", "电机校准专项。"],
            ["静态电流上位机", "静态电流测试专项。"],
            ["imu校准上位机", "IMU 校准专项。"],
            ["屏幕测试上位机", "屏幕测试专项。"],
            ["摄像头测试上位机", "摄像头测试专项。"],
            ["信号测试上位机", "WiFi/BLE 等信号测试专项。"],
            ["老化上位机", "老化测试专项。"],
            ["压感上位机", "压感测试专项。"],
            ["板厂测试上位机", "板厂/PCBA 测试专项。"],
            ["按键测试上位机", "按键测试专项。"],
            ["吸力测试上位机", "吸力测试专项。"],
            ["自由工站上位机", "可编排流程的通用工站（推荐多数新品工序使用）。"],
        ],
        col_widths=[5, 11],
    )

    add_heading_cn(doc, "5.1.3 权限说明", 3)
    add_para(
        doc,
        "「功能设置」仅对管理员或工艺工程师账号显示。若菜单不可见，请使用「帮助 → 切换账号」更换有权限账号。",
        first_line=True,
    )

    add_heading_cn(doc, "5.2 一拖多布局（工位数量）", 2)
    add_steps(
        doc,
        [
            "打开「功能设置」→「关键设置」→「1拖多设置」。",
            "填写「行」「列」（例如行=2、列=2 表示一拖四）。",
            "保存并重新进入工站窗口后，主界面按行列显示多个工位面板。",
            "每个工位独立扫码、独立结果显示；共享外设（如双通道电源）需在步骤参数中配置，见第 8 章。",
        ],
    )

    add_heading_cn(doc, "5.3 自由工站内流程工站（Profile）切换", 2)
    add_para(
        doc,
        "在已选择「自由工站上位机」的前提下，不同产品/工序使用不同的流程工站配置（Profile）。"
        "切换 Profile 不会重启程序，但必须保存流程后才会在开测时生效。",
        first_line=True,
    )
    add_heading_cn(doc, "5.3.1 操作步骤", 3)
    add_steps(
        doc,
        [
            "打开「功能设置」→ 页签「测试流程编排」。",
            "（可选）在「产品名字」中筛选产品。",
            "在「工站」下拉框中选择目标流程工站（例如「Wellness Warm组装吸力测试工站」「M8组装按键测试工站」）。",
            "确认主流程列表与失败执行区内容符合当前工序。",
            "点击「保存流程」。",
            "返回自由工站窗口，确认首 Tab 名称已变为所选流程工站显示名。",
            "按第 7 章连接串口并扫码验证一条完整流程。",
        ],
    )
    add_heading_cn(doc, "5.3.2 工站管理（新建/复制/重命名/删除）", 3)
    add_para(doc, "在「测试流程编排」页通过「工站管理」菜单可进行：", first_line=True)
    add_para(doc, "a）新建工站：创建新的流程工站 Profile；")
    add_para(doc, "b）重命名工站：修改显示名称；")
    add_para(doc, "c）复制工站：基于已有工站复制流程与步骤覆盖；")
    add_para(doc, "d）删除工站：删除不再使用的 Profile（操作前请确认无产线依赖）。")
    add_para(
        doc,
        "云端场景可使用「下载工站用例」「上传本工站用例」与云端草稿同步，按工厂策略执行。",
        first_line=True,
    )

    # ---- 6 步骤配置（重点） ----
    add_heading_cn(doc, "6 工站内部步骤配置操作说明", 1)
    add_para(
        doc,
        "本章说明自由工站「测试流程编排」中功能块的增删改查，以及「测试项配置」对话框各字段含义。"
        "专项工站（按键/吸力等）步骤一般固化在程序内，不以本章方式编排。",
        first_line=True,
    )

    add_heading_cn(doc, "6.1 进入流程编排", 2)
    add_steps(
        doc,
        [
            "菜单「功能设置」→「上位机设置」→「测试流程编排」。",
            "选择要编辑的流程工站（见 5.3）。",
            "上方为主测试流程区；下方为「测试失败执行区域（运行失败时调用）」。",
        ],
    )

    add_heading_cn(doc, "6.2 主流程与失败执行区", 2)
    add_table(
        doc,
        ["区域", "用途"],
        [
            ["主测试流程", "正常测试按顺序执行的功能块列表。"],
            ["测试失败执行区域", "主流程失败后的收尾步骤，如关闭电源、抬气缸、断开连接等。"],
            ["测试失败时结束测试流程", "勾选后：主流程任一步失败则停止主流程并执行失败区。"],
        ],
        col_widths=[5, 11],
    )
    add_para(doc, "必须点击「保存流程」后，拖拽排序与增删块才会写入配置。", first_line=True)

    add_heading_cn(doc, "6.3 添加 / 移除 / 排序功能块", 2)
    add_heading_cn(doc, "6.3.1 添加已有块", 3)
    add_steps(
        doc,
        [
            "在主流程或失败区空白处单击右键，选择「添加已有块」。",
            "在弹出列表中通过「搜索功能块名称」查找目标步骤。",
            "选中后确认添加；块将出现在流程列表中。",
        ],
    )
    add_heading_cn(doc, "6.3.2 添加空白块", 3)
    add_steps(
        doc,
        [
            "右键 →「添加空白块」，创建新步骤。",
            "右键该块 →「打开设置」，按 6.4 完成配置并保存。",
        ],
    )
    add_heading_cn(doc, "6.3.3 排序与移除", 3)
    add_steps(
        doc,
        [
            "拖拽功能块调整顺序，或选中后点击「上移」「下移」。",
            "右键功能块 →「从流程移除」：仅从本流程去掉，不删除磁盘上的步骤库文件。",
            "在「添加已有块」对话框中删除功能块：会删除步骤 ini，需谨慎。",
            "点击「保存流程」。",
        ],
    )
    add_heading_cn(doc, "6.3.4 单步运行（调试）", 3)
    add_steps(
        doc,
        [
            "先打开自由工站窗口并保持运行。",
            "在流程编排中右键目标功能块 →「运行」。",
            "该功能用于工艺调试，一般不过站；调试完成后请用完整扫码流程验证。",
        ],
    )

    add_heading_cn(doc, "6.4 测试项配置（打开设置）", 2)
    add_para(
        doc,
        "对任一功能块：右键 →「打开设置」，进入「测试项配置」对话框。配置完成后点击「Save」保存，再返回流程页点击「保存流程」。",
        first_line=True,
    )

    add_heading_cn(doc, "6.4.1 测试项信息", 3)
    add_table(
        doc,
        ["字段", "说明与操作要点"],
        [
            ["名称", "步骤在界面与结果表中的显示名称。"],
            ["上报MES的字段", "过站字段名；同一流程内勿重复。"],
            ["测试时弹出操作提示", "勾选后测试到该步弹出提示，供操作员确认。"],
            ["提示文字", "弹窗显示的操作说明。"],
            ["不发送指令（PromptOnly）", "勾选后本步仅提示确认，不向设备发送指令。"],
        ],
        col_widths=[5, 11],
    )

    add_heading_cn(doc, "6.4.2 测试指令", 3)
    add_para(doc, "按工序选择通信通道与指令：", first_line=True)
    add_table(
        doc,
        ["字段", "说明"],
        [
            ["通信类型", "产品蓝牙通信、产品串口通信、Dongle通信、云端交互、治具通信、Modbus通信、SCPI通信。"],
            ["产品协议", "当通信类型为产品蓝牙等时选择协议，如 Qfctp、Qpb、Qroot。"],
            ["操作方式", "设置（Set）或读取（Get）。"],
            ["指令内容", "具体设备指令名称（下拉选择）。"],
            ["指令参数", "参数表两列：「参数名」（显示中文）、「参数值」。悬停参数名可查看英文键；点「添加一行」时输入英文参数名。"],
        ],
        col_widths=[4, 12],
    )
    add_para(
        doc,
        "示例：配置程控电源时通信类型选「SCPI通信」，指令选配置源通道类命令，在参数中填写电压、限流、VISA 地址或共享参数（见第 8 章）。",
        first_line=True,
    )

    add_heading_cn(doc, "6.4.3 等待时间", 3)
    add_table(
        doc,
        ["字段", "说明"],
        [
            ["执行前等待", "本步真正执行前延时（毫秒或界面单位以实机为准）。"],
            ["完成后等待", "本步结束后延时再进入下一步。"],
            ["指令超时(毫秒)", "等待设备回包的最长时间。"],
        ],
        col_widths=[4, 12],
    )

    add_heading_cn(doc, "6.4.4 合格判定（Gate）", 3)
    add_steps(
        doc,
        [
            "勾选「根据设备回传结果判定通过/失败」。",
            "选择「回传数据类型」「判定项目」。",
            "选择「判定方式」：在范围内、大于、小于、等于、版本比对等。",
            "填写最小值/最大值或期望文本（按判定方式）。",
            "多字段判定时，在判定表中勾选启用项并分别配置。",
        ],
    )
    add_para(
        doc,
        "说明：部分按键步骤在 Gate 开启且等待设备上报时，弹窗可能不显示「是」按钮，由按键上报自动进入下一步；"
        "若 Gate 关闭且仅为提示，则需操作员点确认继续。",
        first_line=True,
    )

    add_heading_cn(doc, "6.4.5 特殊步骤（Hook）", 3)
    add_steps(
        doc,
        [
            "勾选「不用测试指令，走预置复杂流程」。",
            "在「预置流程类型」中选择 Hook（如蓝牙按名连接、Dongle 吸力相关流程等）。",
            "仅工程师维护复杂逻辑时使用；产线步骤优先使用普通指令+Gate。",
        ],
    )

    add_heading_cn(doc, "6.5 步骤配置推荐流程（检查清单）", 2)
    add_para(doc, "a）明确本步目的：发指令、等回传卡控、仅提示，或 Hook；")
    add_para(doc, "b）选对通信类型与指令，填齐参数；")
    add_para(doc, "c）需要卡控则打开 Gate 并填阈值；")
    add_para(doc, "d）Save 步骤 → 回到编排页保存流程；")
    add_para(doc, "e）可选：右键「运行」单步验证；")
    add_para(doc, "f）自由工站完整扫码跑通主流程与失败收尾。")

    # ---- 7 操作员日常 ----
    add_heading_cn(doc, "7 操作员日常测试操作", 1)
    add_heading_cn(doc, "7.1 开测前检查", 2)
    add_steps(
        doc,
        [
            "确认窗口为当前工序工站；自由工站首 Tab 名称与产品流程一致。",
            "连接所需串口（Dongle/治具/产品串口/万用表）。",
            "按需勾选「是否过站」「是否从mes获取mac」。",
            "扫码枪切换为英文输入。",
        ],
    )
    add_heading_cn(doc, "7.2 连接串口", 2)
    add_para(doc, "a）Dongle：首行串口下拉选口 →「连接串口」；")
    add_para(doc, "b）治具：治具串口区连接，或菜单「连接治具串口」；")
    add_para(doc, "c）产品串口(仪器)、万用表串口：在对应区域选口并连接。")
    add_heading_cn(doc, "7.3 开始与停止", 2)
    add_steps(
        doc,
        [
            "将焦点置于「输入SN获取mac：」或「输入MAC连接：」。",
            "扫描或输入后按回车，软件自动执行流程。",
            "点击「停止测试」可中止；失败时可能自动执行失败区收尾步骤。",
            "在结果表与「日志」Tab 查看通过/失败详情。",
        ],
    )

    # ---- 8 VISA 共享 ----
    add_heading_cn(doc, "8 程控电源与一拖多共享配置（步骤参数）", 1)
    add_para(
        doc,
        "多工位共用程控电源或温度仪时，在步骤「指令参数」中配置，不依赖上位机设置全局开关。",
        first_line=True,
    )
    add_heading_cn(doc, "8.1 共享电源典型参数", 2)
    add_table(
        doc,
        ["界面中文名（示例）", "含义 / 建议"],
        [
            ["启用多工位共享外设", "填 true"],
            ["每台设备对应工位数", "2 表示每 2 工位共一台；3 表示每 3 工位共一台"],
            ["第 1 / 第 2 台程控电源 VISA 地址", "分别对应设备 0、1 的地址"],
            ["选通道 SCPI", "如 Agilent：INST OUT%1"],
            ["电压 (V) / 限流 (A)", "输出设定值"],
        ],
        col_widths=[6, 10],
    )
    add_para(doc, "stationsPerDevice=2 且一拖四时：工位1→设备0通道1，工位2→设备0通道2，工位3→设备1通道1，工位4→设备1通道2。", first_line=True)
    add_heading_cn(doc, "8.2 注意事项", 2)
    add_para(doc, "a）使用 GPIB 前关闭可能独占仪器的软件；")
    add_para(doc, "b）配置步执行后，同流程后续开关/读电流会按工位复用地址与通道；")
    add_para(doc, "c）选通命令以仪器手册为准。")

    # ---- 9 FAQ ----
    add_heading_cn(doc, "9 常见问题", 1)
    add_table(
        doc,
        ["现象", "可能原因", "处理建议"],
        [
            ["看不到功能设置", "账号无权限", "切换管理员/工艺工程师账号"],
            ["切换工站无变化", "未完成重启", "确认程序已重启并核对窗口标题"],
            ["流程不更新", "未点保存流程", "编排后点击「保存流程」"],
            ["扫码无反应", "输入法或焦点不对", "英文输入，点击 SN/MAC 框后再扫"],
            ["串口失败", "占用或选错口", "关闭占用程序，重选 COM"],
            ["两工位抢电源通道", "未配置共享参数", "检查 sharedPair 与地址表"],
            ["单步运行无效", "未开自由工站窗", "先打开自由工站再右键运行"],
        ],
        col_widths=[4, 5, 7],
    )

    # ---- 10 角色 ----
    add_heading_cn(doc, "10 角色分工小结", 1)
    add_table(
        doc,
        ["角色", "推荐操作路径"],
        [
            ["操作员", "登录 → 进工站窗 → 连串口 → 扫 SN/MAC → 看结果/日志"],
            [
                "工艺工程师",
                "功能设置 → 工站类型切换（第5章）→ 流程 Profile → 步骤配置（第6章）→ 可选单步/整流程验证",
            ],
        ],
        col_widths=[3.5, 12.5],
    )

    # ---- 修订记录 ----
    add_heading_cn(doc, "修订记录", 1)
    add_table(
        doc,
        ["版本", "日期", "说明"],
        [
            [
                "A/0",
                "2026-07-27",
                "初稿发布：规范体例；重点增加工站切换（类型/Profile/一拖多）与工站内部步骤配置操作说明。",
            ],
        ],
        col_widths=[2.5, 3.5, 10],
    )
    add_para(doc, "（完）", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print("OK", OUT)


if __name__ == "__main__":
    build()
